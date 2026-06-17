#!/usr/bin/env python3
"""Extract, parse, fill and merge lesson-plan docx files."""

from __future__ import annotations

import argparse
import copy
import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from pypdf import PdfReader

DEFAULT_FONT_NAME = "FangSong"
DEFAULT_FONT_SIZE = 12
STEP_NAMES = ["温故知新", "新课引入", "知识探究", "总结评价"]
DEFAULT_STEP_MINUTES = [8, 7, 50, 15]
STEP_BASE_MINUTES = [5, 5, 20, 5]
STEP_TOTAL_MINUTES = 80

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join([line for line in lines if line]).strip()


def read_txt(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_text(input_file: Path) -> str:
    if not input_file.exists():
        raise FileNotFoundError(f"Input file not found: {input_file}")

    suffix = input_file.suffix.lower()
    if suffix == ".txt":
        return normalize_text(read_txt(input_file))
    if suffix == ".docx":
        doc = Document(str(input_file))
        parts: List[str] = []
        for paragraph in doc.paragraphs:
            text = normalize_text(paragraph.text)
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_parts = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
                if row_parts:
                    parts.append(" | ".join(row_parts))
        return normalize_text("\n".join(parts))
    if suffix == ".pdf":
        reader = PdfReader(str(input_file))
        pages: List[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(page_text)
        return normalize_text("\n".join(pages))

    raise ValueError("Unsupported file type. Only txt/docx/pdf are allowed.")


def apply_run_font(run, font_name: str, font_size: int, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)


def apply_single_line_spacing(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.line_spacing = 1.0
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.left_indent = Pt(0)

    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:hanging"), "0")
    ind.set(qn("w:leftChars"), "0")
    ind.set(qn("w:rightChars"), "0")
    ind.set(qn("w:firstLineChars"), "0")
    ind.set(qn("w:hangingChars"), "0")


def strip_line_indent(text: str) -> str:
    return text.lstrip(" \t\u3000\xa0")


def set_cell_text(
    cell,
    text: str,
    font_name: str,
    font_size: int,
    alignment: WD_PARAGRAPH_ALIGNMENT | None = None,
) -> None:
    content = (text or "").strip()
    lines = [strip_line_indent(line) for line in content.split("\n")] if content else [""]
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        apply_single_line_spacing(p)
        p.alignment = alignment or WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run(line)
        apply_run_font(run, font_name, font_size)


def set_paragraph_text(paragraph, text: str, font_name: str, font_size: int, bold: bool = False) -> None:
    paragraph.text = ""
    apply_single_line_spacing(paragraph)
    run = paragraph.add_run(text)
    apply_run_font(run, font_name, font_size, bold=bold)


def left_align_nonempty_paragraphs(cell) -> None:
    for paragraph in cell.paragraphs:
        if (paragraph.text or "").strip():
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def fix_required_left_alignment(doc: Document) -> None:
    for table in doc.tables:
        if len(table.rows) >= 2 and len(table.rows[1].cells) >= 3:
            label = normalize_text(table.cell(1, 0).text).replace(" ", "")
            if "授课班级" in label:
                left_align_nonempty_paragraphs(table.cell(1, 2))

        if len(table.rows) >= 10 and len(table.rows[6].cells) >= 1:
            step_text = normalize_text(table.cell(6, 0).text).replace(" ", "")
            if "温故知新" in step_text:
                for row in range(6, 10):
                    left_align_nonempty_paragraphs(table.cell(row, 0))


def read_str(data: Dict[str, Any], key: str, default: str = "") -> str:
    value = data.get(key, default)
    return value if isinstance(value, str) else default


def distribute_slots(lengths: List[int], total_slots: int) -> List[int]:
    length_sum = sum(lengths)
    if total_slots <= 0 or length_sum <= 0:
        return [0 for _ in lengths]

    raw_allocations = [(value / length_sum) * total_slots for value in lengths]
    rounded = [int(value) for value in raw_allocations]
    remaining = total_slots - sum(rounded)
    fractions = sorted(
        [(raw_allocations[idx] - rounded[idx], idx) for idx in range(len(lengths))],
        reverse=True,
    )
    for _, idx in fractions[:remaining]:
        rounded[idx] += 1
    return rounded


def compute_step_minutes(steps: List[Dict[str, Any]]) -> List[int]:
    if len(steps) != 4:
        return DEFAULT_STEP_MINUTES[:]

    text_lengths: List[int] = []
    for step in steps:
        content = (step.get("content") or "").strip()
        activity = (step.get("activity") or "").strip()
        intent = (step.get("intent") or "").strip()
        effective_length = len(content) * 2 + len(activity) * 2 + len(intent)
        text_lengths.append(max(1, effective_length))

    base_slots = [max(1, minute // 5) for minute in STEP_BASE_MINUTES]
    total_slots = STEP_TOTAL_MINUTES // 5
    flexible_slots = max(0, total_slots - sum(base_slots))
    extra_slots = distribute_slots(text_lengths, flexible_slots)
    return [(base_slots[idx] + extra_slots[idx]) * 5 for idx in range(4)]


def normalize_steps(raw_steps: Any) -> List[Dict[str, Any]]:
    steps: List[Dict[str, Any]] = []
    if isinstance(raw_steps, list):
        for idx in range(4):
            item = raw_steps[idx] if idx < len(raw_steps) and isinstance(raw_steps[idx], dict) else {}
            steps.append(
                {
                    "name": STEP_NAMES[idx],
                    "minutes": DEFAULT_STEP_MINUTES[idx],
                    "content": item.get("content", "") if isinstance(item.get("content"), str) else "",
                    "activity": item.get("activity", "") if isinstance(item.get("activity"), str) else "",
                    "intent": item.get("intent", "") if isinstance(item.get("intent"), str) else "",
                }
            )
    else:
        for idx in range(4):
            steps.append(
                {
                    "name": STEP_NAMES[idx],
                    "minutes": DEFAULT_STEP_MINUTES[idx],
                    "content": "",
                    "activity": "",
                    "intent": "",
                }
            )

    computed_minutes = compute_step_minutes(steps)
    for idx, minute in enumerate(computed_minutes):
        steps[idx]["minutes"] = minute
    return steps


def replace_lesson_heading(doc: Document, lesson_label: str, font_name: str) -> None:
    target = (lesson_label or "").strip()
    if not target:
        return

    pattern = re.compile(r"^教案\s*1$")
    for paragraph in doc.paragraphs:
        compact_text = normalize_text(paragraph.text).replace(" ", "")
        if pattern.match(compact_text):
            set_paragraph_text(paragraph, target, font_name, 16, bold=True)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            return


def find_row_with_keyword(table, keywords: List[str]) -> int:
    normalized_keywords = [k.replace(" ", "") for k in keywords]
    for row_idx, row in enumerate(table.rows):
        row_text = "".join([cell.text for cell in row.cells]).replace(" ", "")
        for keyword in normalized_keywords:
            if keyword and keyword in row_text:
                return row_idx
    return -1


def is_legacy_process_template(table) -> bool:
    if len(table.rows) < 15 or len(table.columns) < 4:
        return False

    row0 = normalize_text(table.cell(0, 0).text).replace(" ", "")
    row4 = normalize_text(table.cell(4, 0).text).replace(" ", "")
    row10 = normalize_text(table.cell(10, 0).text).replace(" ", "")
    row13 = normalize_text(table.cell(13, 0).text).replace(" ", "")
    return (
        "课前准备" in row0
        and "课中教学" in row4
        and "课后拓展" in row10
        and "考核评价" in row13
    )


def fill_docx(template: Path, structured_json: Path, output_file: Path, font_name: str, font_size: int) -> None:
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")
    if not structured_json.exists():
        raise FileNotFoundError(f"Structured json not found: {structured_json}")

    data = json.loads(structured_json.read_text(encoding="utf-8-sig"))
    if not isinstance(data, dict):
        raise ValueError("Structured json root must be an object.")

    doc = Document(str(template))
    if len(doc.tables) < 2:
        raise ValueError("Template must contain at least two tables.")

    replace_lesson_heading(doc, read_str(data, "lesson_label"), font_name)

    t0 = doc.tables[0]
    t1 = doc.tables[1]
    steps = normalize_steps(data.get("in_class_steps"))

    pre_class = data.get("pre_class", {}) if isinstance(data.get("pre_class"), dict) else {}
    strategy_adjust = data.get("strategy_adjust", {}) if isinstance(data.get("strategy_adjust"), dict) else {}
    post_class = data.get("post_class", {}) if isinstance(data.get("post_class"), dict) else {}

    set_cell_text(t0.cell(0, 2), read_str(data, "topic"), font_name, font_size)
    set_cell_text(t0.cell(0, 8), read_str(data, "reviewer"), font_name, font_size)
    set_cell_text(
        t0.cell(1, 2),
        read_str(data, "class_name"),
        font_name,
        font_size,
        WD_PARAGRAPH_ALIGNMENT.LEFT,
    )
    set_cell_text(t0.cell(1, 4), read_str(data, "date"), font_name, font_size)
    set_cell_text(t0.cell(1, 8), read_str(data, "location"), font_name, font_size)
    set_cell_text(t0.cell(2, 2), read_str(data, "course_type"), font_name, font_size)
    set_cell_text(t0.cell(2, 8), read_str(data, "teaching_mode"), font_name, font_size)

    set_cell_text(t0.cell(3, 2), read_str(data, "xueqing_knowledge"), font_name, font_size)
    set_cell_text(t0.cell(4, 2), read_str(data, "xueqing_mindset"), font_name, font_size)
    set_cell_text(t0.cell(5, 2), read_str(data, "xueqing_cognition"), font_name, font_size)

    set_cell_text(t0.cell(6, 2), read_str(data, "key_points"), font_name, font_size)
    set_cell_text(t0.cell(6, 7), read_str(data, "key_measures"), font_name, font_size)
    set_cell_text(t0.cell(7, 2), read_str(data, "difficult_points"), font_name, font_size)
    set_cell_text(t0.cell(7, 7), read_str(data, "difficult_measures"), font_name, font_size)

    set_cell_text(t0.cell(8, 2), read_str(data, "goal_knowledge"), font_name, font_size)
    set_cell_text(t0.cell(9, 2), read_str(data, "goal_ability"), font_name, font_size)
    set_cell_text(t0.cell(10, 2), read_str(data, "goal_quality"), font_name, font_size)

    set_cell_text(t0.cell(11, 2), read_str(data, "methods"), font_name, font_size)
    set_cell_text(t0.cell(12, 2), read_str(data, "process_tools"), font_name, font_size)
    set_cell_text(t0.cell(13, 2), read_str(data, "resources"), font_name, font_size)

    legacy_process_layout = is_legacy_process_template(t1)

    set_cell_text(t1.cell(2, 1), read_str(pre_class, "content"), font_name, font_size)
    set_cell_text(t1.cell(2, 2), read_str(pre_class, "activity"), font_name, font_size)
    set_cell_text(t1.cell(2, 3), read_str(pre_class, "intent"), font_name, font_size)

    set_cell_text(t1.cell(3, 1), read_str(strategy_adjust, "content"), font_name, font_size)
    set_cell_text(t1.cell(3, 2), read_str(strategy_adjust, "activity"), font_name, font_size)
    set_cell_text(t1.cell(3, 3), read_str(strategy_adjust, "intent"), font_name, font_size)

    for idx, step in enumerate(steps):
        row = 6 + idx
        set_cell_text(
            t1.cell(row, 0),
            f"{step['name']}（{step['minutes']}分钟）",
            font_name,
            font_size,
            WD_PARAGRAPH_ALIGNMENT.LEFT,
        )
        set_cell_text(t1.cell(row, 1), step["content"], font_name, font_size)
        set_cell_text(t1.cell(row, 2), step["activity"], font_name, font_size)
        set_cell_text(t1.cell(row, 3), step["intent"], font_name, font_size)

    if legacy_process_layout:
        set_cell_text(t1.cell(12, 1), read_str(post_class, "content"), font_name, font_size)
        set_cell_text(t1.cell(12, 2), read_str(post_class, "activity"), font_name, font_size)
        set_cell_text(t1.cell(12, 3), read_str(post_class, "intent"), font_name, font_size)
        set_cell_text(t1.cell(14, 0), read_str(data, "evaluation"), font_name, font_size)
    else:
        post_header_row = find_row_with_keyword(t1, ["课后拓展", "课后拓展教学活动", "postclass"])
        post_row = min(post_header_row + 2, len(t1.rows) - 1) if post_header_row >= 0 else min(13, len(t1.rows) - 1)
        set_cell_text(t1.cell(post_row, 1), read_str(post_class, "content"), font_name, font_size)
        set_cell_text(t1.cell(post_row, 2), read_str(post_class, "activity"), font_name, font_size)
        set_cell_text(t1.cell(post_row, 3), read_str(post_class, "intent"), font_name, font_size)

        eval_header_row = find_row_with_keyword(t1, ["考核评价", "evaluation"])
        eval_row = min(eval_header_row + 1, len(t1.rows) - 1) if eval_header_row >= 0 else min(14, len(t1.rows) - 1)
        set_cell_text(t1.cell(eval_row, 0), read_str(data, "evaluation"), font_name, font_size)

    output_file.parent.mkdir(parents=True, exist_ok=True)
    fix_required_left_alignment(doc)
    doc.save(str(output_file))


def parse_preface_meta(paragraphs) -> Dict[str, str]:
    meta = {
        "class_name": "",
        "theory_location": "",
        "practice_location": "",
        "course_type": "",
        "teaching_mode": "",
        "reviewer": "",
    }

    lines = [normalize_text(paragraph.text) for paragraph in paragraphs]
    lines = [line for line in lines if line]

    if len(lines) >= 1:
        meta["class_name"] = lines[0].replace("：", ":").split(":", 1)[-1].strip()
    if len(lines) >= 2:
        meta["theory_location"] = lines[1].replace("：", ":").split(":", 1)[-1].strip()
    if len(lines) >= 3:
        meta["practice_location"] = lines[2].replace("：", ":").split(":", 1)[-1].strip()
    if len(lines) >= 4:
        meta["course_type"] = lines[3].replace("：", ":").split(":", 1)[-1].strip()
    if len(lines) >= 5:
        meta["teaching_mode"] = lines[4].replace("：", ":").split(":", 1)[-1].strip()
    if len(lines) >= 6:
        meta["reviewer"] = lines[5].replace("：", ":").split(":", 1)[-1].strip()

    return meta


def detect_lesson_mode(theory_value: str, practice_value: str) -> str:
    theory_hit = theory_value.strip() == "2"
    practice_hit = practice_value.strip() == "2"
    if theory_hit:
        return "理论课"
    if practice_hit:
        return "实践课"
    return ""


def parse_docx_source(input_file: Path) -> Dict[str, Any]:
    if input_file.suffix.lower() != ".docx":
        raise ValueError("Only docx can be parsed into lessons.")
    if not input_file.exists():
        raise FileNotFoundError(f"Input file not found: {input_file}")

    doc = Document(str(input_file))
    if not doc.tables:
        raise ValueError("Source docx must contain at least one table.")

    meta = parse_preface_meta(doc.paragraphs)
    table = doc.tables[0]
    lessons: List[Dict[str, Any]] = []

    for row in table.rows[2:]:
        cells = row.cells
        if len(cells) < 6:
            continue

        lesson_no = normalize_text(cells[0].text)
        date = normalize_text(cells[1].text)
        topic = normalize_text(cells[2].text)
        content = normalize_text(cells[3].text)
        theory_value = normalize_text(cells[4].text)
        practice_value = normalize_text(cells[5].text)

        if not any([lesson_no, date, topic, content, theory_value, practice_value]):
            continue

        lesson_type = detect_lesson_mode(theory_value, practice_value)
        if not lesson_type:
            continue

        location = meta["theory_location"] if lesson_type == "理论课" else meta["practice_location"]
        display_title = f"教案{lesson_no}" if lesson_no else f"教案{len(lessons) + 1}"
        topic_value = topic or display_title
        content_value = content or topic_value

        lessons.append({
            "lesson_no": lesson_no,
            "date": date,
            "topic": topic_value,
            "content": content_value,
            "teaching_type": lesson_type,
            "location": location,
            "display_title": display_title,
        })

    return {
        "meta": meta,
        "lessons": lessons,
    }


def append_element_before_sectpr(target_doc: Document, element) -> None:
    body = target_doc.element.body
    clone = copy.deepcopy(element)
    if body.sectPr is not None:
        body.insert(len(body) - 1, clone)
    else:
        body.append(clone)


def merge_docx(sections_json: Path, output_file: Path, font_name: str, font_size: int) -> None:
    if not sections_json.exists():
        raise FileNotFoundError(f"Sections json not found: {sections_json}")

    sections = json.loads(sections_json.read_text(encoding="utf-8-sig"))
    if not isinstance(sections, list) or not sections:
        raise ValueError("sections_json must be a non-empty array.")

    merged: Document | None = None
    for index, section in enumerate(sections):
        if not isinstance(section, dict):
            continue

        file_path = Path(str(section.get("file_path", "")))
        if not file_path.exists():
            raise FileNotFoundError(f"Section file not found: {file_path}")

        if merged is None:
            merged = Document(str(file_path))
            continue

        page_break = merged.add_paragraph()
        page_break.add_run().add_break(WD_BREAK.PAGE)

        source_doc = Document(str(file_path))
        for element in source_doc.element.body.iterchildren():
            if element.tag.endswith("sectPr"):
                continue
            append_element_before_sectpr(merged, element)

    if merged is None:
        raise ValueError("No valid sections found.")

    output_file.parent.mkdir(parents=True, exist_ok=True)
    fix_required_left_alignment(merged)
    merged.save(str(output_file))


def build_demo_source_doc(output_file: Path) -> None:
    doc = Document()
    doc.add_paragraph("班级：城D2401")
    doc.add_paragraph("教学方式为理论课时授课地点：K408")
    doc.add_paragraph("教学方式为实践课时授课地点：实训楼301")
    doc.add_paragraph("授课类型：理实一体化课")
    doc.add_paragraph("授课形式：线下")
    doc.add_paragraph("审核人：吴一平")
    doc.add_paragraph("")

    table = doc.add_table(rows=5, cols=6)
    table.style = "Table Grid"
    headers1 = ["教案", "授课日期", "授课章节及课题", "教学内容", "教学方式", "教学方式"]
    headers2 = ["教案", "授课日期", "授课章节及课题", "教学内容", "理论", "实践"]
    for idx, value in enumerate(headers1):
        table.cell(0, idx).text = value
    for idx, value in enumerate(headers2):
        table.cell(1, idx).text = value

    rows = [
        ("01", "03/11", "1.1居住区规划设计基本规定", "居住区规划设计基本要求与选址原则。", "2", ""),
        ("02", "03/13", "1.2居住区的构成", "居住区分级、用地分类以及规划成果要求。", "2", ""),
        ("03", "03/18", "1.3居住区布局组合形式", "围绕片块式、轴线式、围合式等布局形式展开实践任务。", "", "2"),
    ]
    for row_index, values in enumerate(rows, start=2):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = value

    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))


def build_guide_doc(output_file: Path) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("教案生成操作说明")
    apply_run_font(run, DEFAULT_FONT_NAME, 16, bold=True)

    doc.add_paragraph("一、上传文件要求")
    doc.add_paragraph("1. 当前仅支持上传一个 docx 教学文稿。")
    doc.add_paragraph("2. 文稿开头请按段落填写固定信息：班级、理论课时授课地点、实践课时授课地点、授课类型、授课形式、审核人。")
    doc.add_paragraph("3. 文稿后面请使用一张总表，每一行对应一份教案。")
    doc.add_paragraph("4. 系统会将所有有效行生成的教案合并到一个 Word 中。")

    doc.add_paragraph("二、固定段落示例")
    doc.add_paragraph("班级：城D2401")
    doc.add_paragraph("教学方式为理论课时授课地点：K408")
    doc.add_paragraph("教学方式为实践课时授课地点：实训楼301")
    doc.add_paragraph("授课类型：理实一体化课")
    doc.add_paragraph("授课形式：线下")
    doc.add_paragraph("审核人：吴一平")

    doc.add_paragraph("三、表格填写要求")
    doc.add_paragraph("1. 第一张表作为课程清单。")
    doc.add_paragraph("2. 表头应包含：教案、授课日期、授课章节及课题、教学内容、教学方式（理论/实践）。")
    doc.add_paragraph("3. 只有教学方式单元格值等于 2 时，系统才认定该行有效。")
    doc.add_paragraph("4. 理论列和实践列都不是 2 的行会被自动跳过。")
    doc.add_paragraph("5. 若理论列和实践列同时为 2，则按理论课处理。")

    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))


def build_blank_template_doc(output_file: Path) -> None:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("教案空表模板")
    apply_run_font(title_run, DEFAULT_FONT_NAME, 16, bold=True)

    table0 = doc.add_table(rows=14, cols=9)
    table0.style = "Table Grid"

    labels = {
        (0, 0): "授课章节及课题",
        (0, 6): "审核人",
        (1, 0): "授课班级",
        (1, 3): "授课日期",
        (1, 6): "授课地点",
        (2, 0): "授课类型",
        (2, 6): "授课形式",
        (3, 0): "学情分析",
        (3, 1): "知识储备",
        (4, 0): "学情分析",
        (4, 1): "学习习惯",
        (5, 0): "学情分析",
        (5, 1): "认知特点",
        (6, 0): "教学重点",
        (6, 1): "内容",
        (6, 5): "解决措施",
        (7, 0): "教学难点",
        (7, 1): "内容",
        (7, 5): "解决措施",
        (8, 0): "知识目标",
        (8, 1): "目标",
        (9, 0): "能力目标",
        (9, 1): "目标",
        (10, 0): "素质目标",
        (10, 1): "目标",
        (11, 0): "教学方法",
        (12, 0): "教学工具",
        (13, 0): "课程资源",
    }
    for (row, col), text in labels.items():
        set_cell_text(table0.cell(row, col), text, DEFAULT_FONT_NAME, DEFAULT_FONT_SIZE)

    doc.add_paragraph("")

    table1 = doc.add_table(rows=15, cols=4)
    table1.style = "Table Grid"
    process_labels = {
        (0, 0): "教学内容",
        (0, 1): "教学活动",
        (0, 2): "教学工具",
        (0, 3): "设计意图",
        (1, 0): "课前准备",
        (1, 1): "教学内容",
        (1, 2): "教学活动",
        (1, 3): "设计意图",
        (2, 0): "教学内容",
        (3, 0): "教学活动",
        (4, 0): "教学过程",
        (4, 1): "教学内容",
        (4, 2): "教学活动",
        (4, 3): "思政意图",
        (5, 0): "环节",
        (5, 1): "教学内容",
        (5, 2): "教学活动",
        (5, 3): "思政意图",
        (6, 0): "温故知新",
        (7, 0): "新课引入",
        (8, 0): "知识探究",
        (9, 0): "总结评价",
        (10, 0): "课后拓展教学活动",
        (13, 0): "考核评价",
    }
    for (row, col), text in process_labels.items():
        set_cell_text(table1.cell(row, col), text, DEFAULT_FONT_NAME, DEFAULT_FONT_SIZE)

    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))


def cmd_extract(args: argparse.Namespace) -> int:
    text = extract_text(Path(args.input_file))
    payload = {"status": "ok", "text": text}
    if args.print_json:
        print(json.dumps(payload, ensure_ascii=False))
    else:
        print(text)
    return 0


def cmd_parse_docx(args: argparse.Namespace) -> int:
    payload = parse_docx_source(Path(args.input_file))
    payload["status"] = "ok"
    print(json.dumps(payload, ensure_ascii=False))
    return 0


def cmd_fill(args: argparse.Namespace) -> int:
    fill_docx(
        template=Path(args.template),
        structured_json=Path(args.structured_json),
        output_file=Path(args.output_file),
        font_name=args.font_name,
        font_size=args.font_size,
    )
    payload = {"status": "ok", "output_file": args.output_file}
    if args.print_json:
        print(json.dumps(payload, ensure_ascii=False))
    else:
        print(args.output_file)
    return 0


def cmd_merge(args: argparse.Namespace) -> int:
    merge_docx(
        sections_json=Path(args.sections_json),
        output_file=Path(args.output_file),
        font_name=args.font_name,
        font_size=args.font_size,
    )
    payload = {"status": "ok", "output_file": args.output_file}
    if args.print_json:
        print(json.dumps(payload, ensure_ascii=False))
    else:
        print(args.output_file)
    return 0


def cmd_demo_source(args: argparse.Namespace) -> int:
    build_demo_source_doc(Path(args.output_file))
    payload = {"status": "ok", "output_file": args.output_file}
    if args.print_json:
        print(json.dumps(payload, ensure_ascii=False))
    else:
        print(args.output_file)
    return 0


def cmd_guide_doc(args: argparse.Namespace) -> int:
    build_guide_doc(Path(args.output_file))
    payload = {"status": "ok", "output_file": args.output_file}
    if args.print_json:
        print(json.dumps(payload, ensure_ascii=False))
    else:
        print(args.output_file)
    return 0


def cmd_blank_template(args: argparse.Namespace) -> int:
    build_blank_template_doc(Path(args.output_file))
    payload = {"status": "ok", "output_file": args.output_file}
    if args.print_json:
        print(json.dumps(payload, ensure_ascii=False))
    else:
        print(args.output_file)
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Jiaoan text extraction and template filling pipeline.")
    sub = parser.add_subparsers(dest="command", required=True)

    extract = sub.add_parser("extract", help="Extract text from txt/docx/pdf.")
    extract.add_argument("--input-file", required=True, help="Path to txt/docx/pdf file.")
    extract.add_argument("--print-json", action="store_true", help="Print JSON output.")
    extract.set_defaults(func=cmd_extract)

    parse_docx = sub.add_parser("parse-docx", help="Parse meta and lessons from a docx source file.")
    parse_docx.add_argument("--input-file", required=True, help="Path to source docx.")
    parse_docx.add_argument("--print-json", action="store_true", help="Print JSON output.")
    parse_docx.set_defaults(func=cmd_parse_docx)

    fill = sub.add_parser("fill", help="Fill the first form in docx template.")
    fill.add_argument("--template", required=True, help="Path to template docx.")
    fill.add_argument("--structured-json", required=True, help="Path to structured json.")
    fill.add_argument("--output-file", required=True, help="Output docx file path.")
    fill.add_argument("--font-name", default=DEFAULT_FONT_NAME, help="Font name for filled content.")
    fill.add_argument("--font-size", type=int, default=DEFAULT_FONT_SIZE, help="Font size in pt.")
    fill.add_argument("--print-json", action="store_true", help="Print JSON output.")
    fill.set_defaults(func=cmd_fill)

    merge = sub.add_parser("merge", help="Merge multiple lesson-plan docx files into one output.")
    merge.add_argument("--sections-json", required=True, help="Path to sections json.")
    merge.add_argument("--output-file", required=True, help="Output docx file path.")
    merge.add_argument("--font-name", default=DEFAULT_FONT_NAME, help="Font name for inserted titles.")
    merge.add_argument("--font-size", type=int, default=DEFAULT_FONT_SIZE, help="Font size in pt.")
    merge.add_argument("--print-json", action="store_true", help="Print JSON output.")
    merge.set_defaults(func=cmd_merge)

    demo_source = sub.add_parser("demo-source", help="Generate a sample source docx.")
    demo_source.add_argument("--output-file", required=True, help="Output docx file path.")
    demo_source.add_argument("--print-json", action="store_true", help="Print JSON output.")
    demo_source.set_defaults(func=cmd_demo_source)

    guide_doc = sub.add_parser("guide-doc", help="Generate an operation guide docx.")
    guide_doc.add_argument("--output-file", required=True, help="Output docx file path.")
    guide_doc.add_argument("--print-json", action="store_true", help="Print JSON output.")
    guide_doc.set_defaults(func=cmd_guide_doc)

    blank_template = sub.add_parser("blank-template", help="Generate a blank lesson-plan template docx.")
    blank_template.add_argument("--output-file", required=True, help="Output docx file path.")
    blank_template.add_argument("--print-json", action="store_true", help="Print JSON output.")
    blank_template.set_defaults(func=cmd_blank_template)

    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    try:
        return args.func(args)
    except Exception as exc:  # pylint: disable=broad-except
        if getattr(args, "print_json", False):
            print(json.dumps({"status": "error", "message": str(exc)}, ensure_ascii=False))
        else:
            print(str(exc), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
